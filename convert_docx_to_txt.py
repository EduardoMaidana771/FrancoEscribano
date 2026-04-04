# Convierte todos los .docx y .doc de la carpeta de Documentos de ejemplos
# a archivos .txt organizados en la misma estructura de carpetas.

import os
import sys
from pathlib import Path

from docx import Document

SOURCE_DIR = Path(r"C:\Users\manue\Downloads\Franco\Documentos de ejemplos")
OUTPUT_DIR = Path(r"C:\Users\manue\source\repos\FrancoEscribano\docs_txt")

stats = {"ok": 0, "skip_temp": 0, "fail": 0, "doc_skip": 0}
failures = []


def convert_docx(src: Path, dst: Path):
    """Lee un .docx y guarda todo su texto en un .txt."""
    doc = Document(str(src))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # También extraer texto de tablas (algunas escrituras usan tablas)
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                lines.append(" | ".join(cells_text))
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n".join(lines), encoding="utf-8")


def main():
    print(f"Origen: {SOURCE_DIR}")
    print(f"Destino: {OUTPUT_DIR}")
    print()

    if not SOURCE_DIR.exists():
        print("ERROR: Carpeta de origen no existe")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    all_files = list(SOURCE_DIR.rglob("*"))
    doc_files = [
        f for f in all_files
        if f.is_file()
        and f.suffix.lower() in (".docx", ".doc")
        and not f.name.startswith("~$")
    ]

    print(f"Archivos encontrados: {len(doc_files)}")
    print()

    for i, src in enumerate(sorted(doc_files), 1):
        rel = src.relative_to(SOURCE_DIR)
        dst = OUTPUT_DIR / rel.with_suffix(".txt")

        if src.suffix.lower() == ".doc":
            # python-docx no soporta .doc antiguo
            stats["doc_skip"] += 1
            print(f"  [{i}/{len(doc_files)}] SKIP (.doc antiguo): {rel}")
            # Intentar como .docx de todos modos (a veces están mal renombrados)
            try:
                convert_docx(src, dst)
                stats["ok"] += 1
                stats["doc_skip"] -= 1
                print(f"    -> Convertido exitosamente como .docx!")
            except Exception:
                failures.append(("doc_format", str(rel)))
            continue

        try:
            convert_docx(src, dst)
            stats["ok"] += 1
            if i % 50 == 0 or i == len(doc_files):
                print(f"  [{i}/{len(doc_files)}] OK: {rel}")
        except Exception as e:
            stats["fail"] += 1
            failures.append(("error", str(rel), str(e)[:100]))
            print(f"  [{i}/{len(doc_files)}] FAIL: {rel} -> {e}")

    print()
    print("=" * 60)
    print(f"RESULTADO:")
    print(f"  Convertidos OK:     {stats['ok']}")
    print(f"  .doc omitidos:      {stats['doc_skip']}")
    print(f"  Fallidos:           {stats['fail']}")
    print(f"  Total procesados:   {len(doc_files)}")
    print()

    if failures:
        print("ARCHIVOS CON PROBLEMAS:")
        for f in failures:
            print(f"  - {f}")

    # Verificar que los .txt se pueden leer
    print()
    print("Verificando lectura de archivos generados...")
    txt_files = list(OUTPUT_DIR.rglob("*.txt"))
    read_ok = 0
    read_fail = 0
    empty = 0
    for tf in txt_files:
        try:
            content = tf.read_text(encoding="utf-8")
            if len(content.strip()) == 0:
                empty += 1
            else:
                read_ok += 1
        except Exception:
            read_fail += 1

    print(f"  Legibles:           {read_ok}")
    print(f"  Vacíos:             {empty}")
    print(f"  Error de lectura:   {read_fail}")
    print(f"  Total .txt:         {len(txt_files)}")


if __name__ == "__main__":
    main()
