"""Rebuild compraventa.docx template from reference document.

Creates a Docxtemplater-compatible template with split runs for
bold/underline control at boundary level, minimal placeholders for
dynamic data, and static legal text embedded in template.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"c:\Users\manue\source\repos\FrancoEscribano")
REFERENCE_DOCX = Path(r"c:\Users\manue\Downloads\Franco\compraventa ultima.docx")
OUTPUT_DOCX = ROOT / "app" / "templates" / "compraventa.docx"

# Font sizes by paragraph range (0-based index)
SIZE_11_5 = Pt(11.5)  # P0-P13
SIZE_12_5 = Pt(12.5)  # P14-P16
SIZE_13 = Pt(13)       # P31


def clear_runs(paragraph):
    """Remove all w:r elements from a paragraph, preserving paragraph properties."""
    p_elem = paragraph._p
    for child in list(p_elem):
        if child.tag == qn("w:r"):
            p_elem.remove(child)


def add_run(paragraph, text, *, bold=False, underline=False, font_size=SIZE_11_5):
    """Add a run with explicit formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.name = "Aptos"
    run.font.size = font_size
    return run


def rebuild_paragraph(paragraph, runs_spec, font_size):
    """Clear a paragraph and rebuild with specified runs.

    runs_spec: list of (text, bold, underline) tuples.
    """
    clear_runs(paragraph)
    for text, bold, ul in runs_spec:
        add_run(paragraph, text, bold=bold, underline=ul, font_size=font_size)


# ── Run specifications per paragraph (0-based) ──────────────────────
# Each entry: (text, bold, underline)
# B = bold only, BU = bold + underline, N = normal

PARAGRAPHS_11_5 = {
    # P2: Intro
    2: [
        ("En la ciudad de {ciudad}, el día {fecha_letras}, las partes que se expresan han convenido celebrar el siguiente contrato de compraventa:", False, False),
    ],
    # P3: 1. PARTE VENDEDORA
    3: [
        ("{vendedora_titulo}", True, False),
        ("{vendedora_texto}", False, False),
    ],
    # P4: 2. PARTE COMPRADORA
    4: [
        ("{compradora_titulo}", True, False),
        ("{compradora_texto}", False, False),
    ],
    # P5: 3. BIEN QUE SE VENDE
    5: [
        ("3. BIEN QUE SE VENDE. - ", True, False),
        ("La parte vendedora vende a la parte compradora y ésta adquiere, ", False, False),
        ("libre de toda obligación, gravamen, multas, infracciones y demás responsabilidades civiles, tributarias y penales, ", True, False),
        ("el siguiente bien: (auto, camión, etc): ", False, False),
        ("AUTOMOTOR:", True, False),
    ],
    # P6: 4. DERECHO QUE SE TRASMITE
    6: [
        ("4. DERECHO QUE SE TRASMITE. (CUOTA PARTE).- ", True, False),
        ("La parte vendedora trasmite a la parte compradora: ", False, False),
        ("LA PROPIEDAD ", True, False),
        ("del vehículo a que se refiere la cláusula 3, a saber las {cuota_parte} avas partes.-", False, False),
    ],
    # P7: 5. PRECIO
    7: [
        ("5. PRECIO. - ", True, False),
        ("El precio de esta compraventa asciende a la suma de {precio_moneda} ", False, False),
        ("{precio_monto}", True, False),
        ("{precio_pago}", False, False),
    ],
    # P8: 6. TRADICIÓN
    8: [
        ("6. TRADICIÓN. - ", True, False),
        ("Como tradición la parte vendedora trasmite a la parte compradora todos los derechos de propiedad y posesión que sobre el referido vehículo le corresponden, el que toma en este acto. -", False, False),
    ],
    # P9: 7. TÍTULO ANTERIOR
    9: [
        ("7. TÍTULO ANTERIOR. – ", True, False),
        ("Se controlará por certificación notarial.", False, False),
    ],
    # P10: 8. DOMICILIO
    10: [
        ("8. DOMICILIO.- ", True, False),
        ("Las partes fijan a los efectos de este contrato y de las actuaciones registrales que se cumplan en el Registro de Vehículos Automotores, los indicados como suyos respectivamente, en la comparecencia. -", False, False),
    ],
    # P11: DECLARACIONES
    11: [
        (". DECLARACIONES. - ", True, False),
        ("La parte vendedora: declara bajo juramento: {declaracion_bps} La parte adquirente: i) Conoce que es su responsabilidad verificar que número de motor y chasis coincidan con el indicado en la libreta de circulación, exonerando de responsabilidad al Escribano actuante;", False, False),
    ],
    # P12: Responsabilidad (no bold, fully dynamic)
    12: [
        ("{declaracion_responsabilidad_texto}", False, False),
    ],
    # P13: 10. SOLICITUD DE INTERVENCIÓN NOTARIAL
    13: [
        ("10. SOLICITUD DE INTERVENCIÓN NOTARIAL. - ", True, False),
        ("Las partes solicitan al Escribano {notary_name}, que certifique el otorgamiento y suscripción de este contrato. - De conformidad las partes suscriben el mismo escrito en el lugar y fecha indicados ut supra. ", False, False),
        ("{firmas_texto}", True, False),
    ],
}

PARAGRAPHS_12_5 = {
    # P14: CERTIFICO QUE
    14: [
        ("CERTIFICO QUE: ", True, True),   # BU
        ("I) Las firmas que anteceden son auténticas, fueron puestas en mi presencia y pertenecen a personas hábiles que no conozco, pero acreditaron su identidad con los documentos respectivos señores ", False, False),
        ("{certifico_nombres}", True, False),   # B
        ("{certifico_cedulas_texto}{certifico_secciones}", False, False),
        (" EN FE DE ELLO ", True, False),   # B
        ("a solicitud de la parte interesada y para su presentación ante quien corresponda expido el presente que sello, signo y firmo en {ciudad} el día {fecha_letras}. ", False, False),
        ("{notary_firma}", True, False),   # B
    ],
    # P15: Protocolización
    15: [
        ("Nº {matriz_numero} Protocolización Preceptiva De Compraventa Automotor. ", True, True),  # BU
        ("{proto_partes} ", True, False),   # B
        ("En {ciudad} el día {fecha_letras} cumpliendo con lo dispuesto por el artículo 292 de la ley 18362 incorporo a mi Registro de Protocolizaciones compraventa automotor de padrón ", False, False),
        ("{vehicle_padron} ", True, False),   # B
        ("con certificación y la presente con el número {matriz_numero}, extendida de folio {folio_start} a {folio_end_display}. ", False, False),
        ("{notary_firma_short}", True, False),   # B
    ],
    # P16: ES PRIMER TESTIMONIO
    16: [
        ("ES PRIMER TESTIMONIO ", True, True),   # BU
        ("que he compulsado de la protocolización que incorpore a mi Registro con el número {matriz_numero} en hoja de papel notarial serie {paper_series_proto} número {paper_number_proto}. ", False, False),
        ("EN FE DE ELLO ", True, False),   # B
        ("y para la parte compradora expido el presente que sello signo y firmo en {ciudad} el día {fecha_letras} en hojas de papel notarial serie {paper_series_testimony} número {paper_numbers_testimony}.", False, False),
    ],
}

# P31: Insurance certificate (conditional via Docxtemplater {#has_insurance})
PARAGRAPH_31 = [
    ("{#has_insurance}CERTIFICO QUE: ", True, True),   # BU
    ("Tuve a la vista Certificado de Seguro Obligatorio póliza número {insurance_policy}, expedido por {insurance_company} y vencimiento el {insurance_expiry}, dando cumplimento a la ley 18412. ", False, False),
    ("EN FE DE ELLO ", True, False),   # B
    ("a solicitud de parte interesada y para su presentación ante quien corresponda expido el presente que sello, signo y firmo en {ciudad} el día {fecha_letras}.{/has_insurance}", False, False),
]

# Table cells: (row, col, text, bold)
TABLE_CELLS = [
    (0, 0, "TIPO: {vehicle_type}", False),
    (0, 1, "MARCA: {vehicle_brand}", False),
    (1, 0, "MODELO: {vehicle_model}", False),
    (1, 1, "AÑO: {vehicle_year}", False),
    (2, 0, "PADRON: {vehicle_padron} de {vehicle_padron_department}", True),
    (2, 1, "MATRICULA: {vehicle_plate}", True),
    (3, 0, "MOTOR: {vehicle_motor}", False),
    (3, 1, "A (combustible): {vehicle_fuel}", False),
    (4, 0, "CHASIS: {vehicle_chassis}", False),
    (4, 1, "CILINDROS: {vehicle_cylinders}", False),
]


def main() -> None:
    if not REFERENCE_DOCX.exists():
        raise FileNotFoundError(f"No existe referencia: {REFERENCE_DOCX}")

    doc = Document(str(REFERENCE_DOCX))

    # P0 and P1 (header): make runs explicitly bold, preserve alignment
    for pi in (0, 1):
        for run in doc.paragraphs[pi].runs:
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = SIZE_11_5

    # P2-P13: 11.5pt font, split runs at bold boundaries
    for pi, runs_spec in PARAGRAPHS_11_5.items():
        rebuild_paragraph(doc.paragraphs[pi], runs_spec, SIZE_11_5)

    # P14-P16: 12.5pt font
    for pi, runs_spec in PARAGRAPHS_12_5.items():
        rebuild_paragraph(doc.paragraphs[pi], runs_spec, SIZE_12_5)

    # P31: 13pt font, conditional insurance certificate
    if len(doc.paragraphs) > 31:
        rebuild_paragraph(doc.paragraphs[31], PARAGRAPH_31, SIZE_13)

    # Vehicle data table
    if not doc.tables:
        raise RuntimeError("La referencia no contiene tabla de datos del vehículo")

    table = doc.tables[0]
    for row, col, value, bold in TABLE_CELLS:
        cell = table.cell(row, col)
        # Clear existing cell content
        for p in cell.paragraphs:
            clear_runs(p)
        p = cell.paragraphs[0]
        add_run(p, value, bold=bold, font_size=SIZE_11_5)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"Template generado: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
